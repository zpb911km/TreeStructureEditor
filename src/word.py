from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def saveAsDocx(text, PATH) -> None:
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    section = doc.sections[0]
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = Pt(10)
    for line in text.split('\n'):
        para.add_run(line + '\r')
    doc.save(PATH)


def readFromDocx(PATH) -> str:
    doc = Document(PATH)
    t = [par.text for par in doc.paragraphs]
    t = t[0]
    print(t)


if __name__ == '__main__':
    PATH = "D:\新建文件夹\Doc1.docx"
    text = '''>英语
    ├>连接词
    │ ├>Ⅰ
    │ │ ├>firstly
    │ │ └>first of all
    │ ├>Ⅱ
    │ │ ├>moreover
    │ │ ├>furthermore
    │ │ ├>better yet
    │ │ ├>what's more
    │ │ ├>additionally
    │ │ ├>in addition
    │ │ ├>Another key point(另一个关键点)
    │ │ └>Equally important(同样重要的是)
    │ ├>Ⅲ
    │ │ ├>most importantly
    │ │ ├>eventually
    │ │ ├>finally
    │ │ ├>all in all
    │ │ ├>without doubtIn conclusion(总之)
    │ │ └>In summary(总之)
    │ ├>转折
    │ │ ├>actually
    │ │ ├>however
    │ │ ├>Nonetheless(尽管如此)
    │ │ ├>In spite of(尽管)
    │ │ ├>Despite(尽管)'''

    saveAsDocx(text, PATH)
    readFromDocx(PATH)
